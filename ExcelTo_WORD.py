import openpyxl
from num_to_word import num_to_word
import docx
import os
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from openpyxl import load_workbook
import docxcompose
import docxcompose.composer
from docxcompose.composer import Composer
from docx import Document as Document_compose

def exc_to_word(w,q,d):
        global doc
        doc=docx.Document(w)
        naam= "uke %&"
        ganv="xzke %&"
        gata="xkVk la0 %&"
        check_para=doc.paragraphs[9]
        global wb
        wb= openpyxl.load_workbook(q,data_only=True)
        #global sheet
        #sheet= wb.active        
        for i in range(0,len(wb.worksheets)):
                sheetname=wb.worksheets[i]
                global curr_row
                curr_row=row-1
                global rows
                rows = sheetname.max_row-1
                files_list=[]
                excel_gaata_p="uu"
                excel_vill_p="ii"
                while curr_row<(rows):                        
                        doc=docx.Document(w)
                        check_para=doc.paragraphs[9]
                        curr_row+=1
                        #name fetching,vill fetching,gata no. fetching
                        excel_name = str(sheetname.cell(curr_row, namecol).value)
                        excel_vill=str(sheetname.cell(curr_row,villcol).value)
                        excel_gaata=str(sheetname.cell(curr_row, gaatacol).value)
                        
                        amount=str(sheetname.cell(curr_row,amountcol).value)
                        s=str(sheetname)
                        global sheetname1
                        sheetname1=(s[12:-2])
                        if excel_name=="None":
                            continue
                        if amount=="None" or len(amount)>8:
                            continue
                        
                        if excel_gaata=="None":
                            excel_gaata=excel_gaata_p
                            print(excel_gaata+"recalling gaata")
                            gaata=gata+" "+excel_gaata
                        if excel_gaata != "None":
                            gaata=gata+" "+excel_gaata
                            excel_gaata_p=excel_gaata
                            print(excel_gaata_p+"updating gaata")
                        if excel_vill=="None":
                            excel_vill=excel_vill_p
                            print(excel_vill+"recalling vill")
                            gaav=ganv+" "+excel_vill
                        if excel_vill != "None":
                            gaav=ganv+" "+excel_vill
                            excel_vill_p=excel_vill
                            print(excel_vill_p+"updating vill")
                        
                        name=naam+" "+excel_name
                        #gaav=ganv+" "+excel_vill
                        #gaata=gata+" "+excel_gaata
                        #amount to word conversion                    
                        print(amount)
                        words=num_to_word(amount, lang='hi',separator=' ')
                        print(words)
                        words1=words+" ek="
                        Dictionary = {"mmmm": amount, "nnnn":words1, "pppp":"D"}
                        #font.bold=True
                        for i in Dictionary:
                            for p in doc.paragraphs:
                                if p.text.find(i)>=0:
                                    p.text=p.text.replace(i,Dictionary[i])
                        #assiging name vill gata into wordfile
                        first = check_para.insert_paragraph_before(name)
                        second=check_para.insert_paragraph_before(gaav)
                        third=check_para.insert_paragraph_before(gaata)
                        # applying font and size
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Kruti Dev 010'
                        font.size = Pt(18)
                        first.style=doc.styles['Normal']
                        second.style=doc.styles['Normal']
                        third.style=doc.styles['Normal']
                        global directory                       
                        first_file_name=str(d)+"/"+str(sheetname1)+str(row)+".docx"
                        directory=str(d)+"/"+str(sheetname1)+str(curr_row)+".docx"
                        doc.save(directory)
                        #files_list.append(directory)                        
                        print("Source path renamed to destination path successfully.")
                        print(name)
                        print("files list appended success")
                        files_list.append(str(directory))
                        print(files_list[0])
                files_list.pop(0)
                combine_all_docx(str(d)+"/"+str(sheetname1)+str(row)+".docx",files_list)
                ###########################
                messagebox.showinfo("Attention!!", "Data moved Succesfully for - "+str(sheetname1)+"."+"\nNow program will delete single files"+
                                    "\nCheck Output Directory")
                #remove leftovers
                curr_row=row-1
                rows = sheetname.max_row-1
                while curr_row<(rows+1):
                    curr_row+=1
                    amount=str(sheetname.cell(curr_row,amountcol).value)
                    excel_name = str(sheetname.cell(curr_row, namecol).value)
                    if excel_name=="None":
                            continue
                    if amount=="None" or len(amount)>8:
                            continue
                    os.remove(str(d)+"/"+str(sheetname1)+str(curr_row)+".docx")
                    print("file deleted "+str(d)+"/"+str(sheetname1)+str(curr_row)+".docx")
        messagebox.showinfo("Attention!!", "Data moved Succesfully for All Sheets!! \n DONE!!!!!")
def combine_all_docx(first_file_name,files_list):
    number_of_sections=len(files_list)
    master = Document_compose(first_file_name)
    composer = Composer(master)
    master.add_page_break()
    for i in range(0, number_of_sections):
         doc_temp = Document_compose(files_list[i])
         if i<number_of_sections-1:
             doc_temp.add_page_break()             
         
         composer.append(doc_temp)
    composer.save(str(dest)+"/"+str(sheetname1)+" All_Merged_In_One"+".docx")
def browsewordFile():
    global wordfilename
    wordfilename= filedialog.askopenfilename(initialdir = "/",title = "Select a File",filetypes = (("Word files","*.docx*"),("all files","*.*")))
    print(wordfilename)
    label_file_explorer1.configure(text="File Opened: "+wordfilename)
def browseexcelFile():
    global excelfilename
    excelfilename= filedialog.askopenfilename(initialdir = "/",title = "Select a File",filetypes = (("Excel files","*.xlsx*"),("all files","*.*")))
    print(excelfilename)
    label_file_explorer2.configure(text="File Opened: "+excelfilename)
def output_dir():
    global dest
    dest = filedialog.askdirectory()    
    label_file_explorer3.configure(text="Output Directory is: "+dest)
    print(dest)
def submit():
    global namecol,villcol,amountcol,gaatacol,row
    namecol=int(col_name.get())
    villcol=int(col_vill.get())
    amountcol=int(col_amount.get())
    gaatacol=int(col_gaata.get())
    row=int(main_row.get())
    messagebox.showinfo("Attention!!", "Values loaded Successfully!")
                                                                                
                                          
window = Tk()
window.title('Excel_to_Word_HOMEMADE')
window.geometry("680x450")
window['background']='#B0E2FF'
window.iconbitmap("newt.ico")


col_name=tk.StringVar()
col_vill=tk.StringVar()
col_amount=tk.StringVar()
col_gaata=tk.StringVar()
main_row=tk.StringVar()
  
# Create a File Explorer label
label_file_explorer1 = Label(window,text = "*Word File",width = 80, height = 1,fg = "black",bg="LightSkyBlue2")
label_file_explorer2 = Label(window,text = "*Excel File",width = 80, height = 1,fg = "black",bg="LightSkyBlue2")
label_file_explorer3 = Label(window,text = "*Output Directory",width =50, height = 1,fg = "black",bg="LightSkyBlue2")
name_label1 = Label(window, text = 'Column of Name', font=('calibre',10, 'bold'),bg="LightSkyBlue1")
name_entry1 = Entry(window,textvariable = col_name, font=('calibre',10,'normal'),width=5)
name_label2 = Label(window, text = 'Column of village', font=('calibre',10, 'bold'),bg="LightSkyBlue1")
name_entry2 = Entry(window,textvariable = col_vill, font=('calibre',10,'normal'),width=5)
name_label3 = Label(window, text = 'Column of Amount', font=('calibre',10, 'bold'),bg="LightSkyBlue1")
name_entry3 = Entry(window,textvariable = col_amount, font=('calibre',10,'normal'),width=5)
name_label4 = Label(window, text = 'Column of Gaata', font=('calibre',10, 'bold'),bg="LightSkyBlue1")
name_entry4 = Entry(window,textvariable = col_gaata, font=('calibre',10,'normal'),width=5)
name_label5 = Label(window, text = "ROW number\n(it should be common \n for all columns)",font=('calibre',10, 'bold'),bg="LightSkyBlue1")
name_entry5 = Entry(window,textvariable = main_row, font=('calibre',10,'normal'),width=5)
sub_btn=Button(window,text = 'Submit Me First', command = submit,borderwidth=3,fg="black",bg="DarkOrange1")
button_explore1 = Button(window, 
                        text = "Browse Word File",borderwidth=5,fg="white",bg="blue",
                        command = lambda:[browsewordFile()])
button_explore2 = Button(window, 
                        text = "Browse Excel File",borderwidth=5,fg="black",bg="orange",
                        command = lambda:[browseexcelFile()])
button_explore3 = Button(window, 
                        text = "Select Output Directory",bg="dark green",fg="white",
                        borderwidth=5,command = lambda:[output_dir()])

  
button_exit = Button(window, 
                     text = "MAGIC BUTTON",width=20,height=2,bg="red",fg="white",borderwidth=8,
                     command = lambda:[exc_to_word(wordfilename,excelfilename,dest)]) 
  
# Grid method is chosen for placing
# the widgets at respective positions 
# in a table like structure by
# specifying rows and columns
label_file_explorer1.grid(column = 4, row = 1)
label_file_explorer2.grid(column =4, row = 3)
label_file_explorer3.grid(column = 4, row = 6)
button_explore1.grid(column = 4, row = 2)
button_explore2.grid(column = 4, row = 4)
button_explore3.grid(column = 4, row = 9)
button_exit.grid(column = 4,row = 11)
name_label1.grid(row=0,column=0)
name_entry1.grid(row=0,column=1)
name_label2.grid(row=1,column=0)
name_entry2.grid(row=1,column=1)
name_label3.grid(row=2,column=0)
name_entry3.grid(row=2,column=1)
name_label4.grid(row=3,column=0)
name_entry4.grid(row=3,column=1)
name_label5.grid(row=5,column=0)
name_entry5.grid(row=5,column=1)
sub_btn.grid(row=7,column=0)

# Let the window wait for any events
window.mainloop()

                                                        

   
    








